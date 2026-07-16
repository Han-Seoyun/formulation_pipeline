"""
GHS 라벨 파이프라인 현황 보고서 생성
출력: reports/GHS_Label_Pipeline_Report_20260714.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path
import copy

BASE_DIR = Path(__file__).parent.parent
OUT_DIR  = BASE_DIR / "reports"
OUT_DIR.mkdir(exist_ok=True)
OUT_PATH = OUT_DIR / "GHS_Label_Pipeline_Report_20260714.docx"

# ── 색상 팔레트 ──────────────────────────────────────────────────────────────
C_NAVY    = RGBColor(0x1F, 0x3A, 0x5F)
C_BLUE    = RGBColor(0x2E, 0x74, 0xB5)
C_LIGHT   = RGBColor(0xD6, 0xE4, 0xF0)
C_RED     = RGBColor(0xC0, 0x39, 0x2B)
C_ORANGE  = RGBColor(0xE6, 0x7E, 0x22)
C_GREEN   = RGBColor(0x1E, 0x8B, 0x4C)
C_GRAY    = RGBColor(0x60, 0x60, 0x60)
C_LGRAY   = RGBColor(0xF2, 0xF2, 0xF2)
C_WHITE   = RGBColor(0xFF, 0xFF, 0xFF)
C_BLACK   = RGBColor(0x00, 0x00, 0x00)

def set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    hex_color = f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), kwargs.get('val', 'single'))
        border.set(qn('w:sz'), kwargs.get('sz', '4'))
        border.set(qn('w:color'), kwargs.get('color', '000000'))
        tcBorders.append(border)
    tcPr.append(tcBorders)

def add_heading(doc, text, level=1, color=None):
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    run = p.add_run(text)
    if level == 1:
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = color or C_NAVY
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)
        # 하단 테두리
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '8')
        bottom.set(qn('w:color'), f"{C_BLUE[0]:02X}{C_BLUE[1]:02X}{C_BLUE[2]:02X}")
        pBdr.append(bottom)
        pPr.append(pBdr)
    elif level == 2:
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = color or C_BLUE
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
    elif level == 3:
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = color or C_NAVY
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
    return p

def add_body(doc, text, bold=False, color=None, indent=False):
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    if indent:
        p.paragraph_format.left_indent = Cm(0.8)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(3)
    return p

def add_bullet(doc, text, level=0, color=None, bold=False):
    p = doc.add_paragraph(style='List Bullet')
    indent = Cm(0.5 + level * 0.6)
    p.paragraph_format.left_indent = indent
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(2)
    return p

def add_kv_table(doc, rows, col_widths=(5, 9)):
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = 'Table Grid'
    for i, (k, v, *rest) in enumerate(rows):
        highlight = rest[0] if rest else None
        k_cell = table.cell(i, 0)
        v_cell = table.cell(i, 1)
        k_cell.width = Cm(col_widths[0])
        v_cell.width = Cm(col_widths[1])
        set_cell_bg(k_cell, C_LGRAY)
        k_run = k_cell.paragraphs[0].add_run(k)
        k_run.font.size = Pt(9.5)
        k_run.font.bold = True
        k_run.font.color.rgb = C_NAVY
        if highlight == 'red':
            set_cell_bg(v_cell, RGBColor(0xFF, 0xEB, 0xEB))
            v_run = v_cell.paragraphs[0].add_run(v)
            v_run.font.color.rgb = C_RED
        elif highlight == 'green':
            set_cell_bg(v_cell, RGBColor(0xE8, 0xF8, 0xEE))
            v_run = v_cell.paragraphs[0].add_run(v)
            v_run.font.color.rgb = C_GREEN
        elif highlight == 'orange':
            set_cell_bg(v_cell, RGBColor(0xFF, 0xF3, 0xE0))
            v_run = v_cell.paragraphs[0].add_run(v)
            v_run.font.color.rgb = C_ORANGE
        else:
            v_run = v_cell.paragraphs[0].add_run(v)
            v_run.font.color.rgb = C_BLACK
        v_run.font.size = Pt(9.5)
        for cell in [k_cell, v_cell]:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.paragraphs[0].paragraph_format.space_before = Pt(3)
            cell.paragraphs[0].paragraph_format.space_after = Pt(3)
    return table

def add_section_table(doc, headers, data_rows, col_widths=None):
    table = doc.add_table(rows=1 + len(data_rows), cols=len(headers))
    table.style = 'Table Grid'
    # 헤더
    hdr_row = table.rows[0]
    for j, h in enumerate(headers):
        cell = hdr_row.cells[j]
        set_cell_bg(cell, C_NAVY)
        run = cell.paragraphs[0].add_run(h)
        run.font.size = Pt(9.5)
        run.font.bold = True
        run.font.color.rgb = C_WHITE
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.paragraphs[0].paragraph_format.space_before = Pt(3)
        cell.paragraphs[0].paragraph_format.space_after = Pt(3)
    # 데이터
    for i, row_data in enumerate(data_rows):
        row = table.rows[i + 1]
        bg = C_LGRAY if i % 2 == 0 else C_WHITE
        for j, (val, *opts) in enumerate(row_data):
            cell = row.cells[j]
            color_hint = opts[0] if opts else None
            if color_hint == 'red':
                set_cell_bg(cell, RGBColor(0xFF, 0xEB, 0xEB))
                run = cell.paragraphs[0].add_run(val)
                run.font.color.rgb = C_RED
            elif color_hint == 'green':
                set_cell_bg(cell, RGBColor(0xE8, 0xF8, 0xEE))
                run = cell.paragraphs[0].add_run(val)
                run.font.color.rgb = C_GREEN
            elif color_hint == 'orange':
                set_cell_bg(cell, RGBColor(0xFF, 0xF3, 0xE0))
                run = cell.paragraphs[0].add_run(val)
                run.font.color.rgb = C_ORANGE
            else:
                set_cell_bg(cell, bg)
                run = cell.paragraphs[0].add_run(val)
                run.font.color.rgb = C_BLACK
            run.font.size = Pt(9.5)
            if col_widths:
                cell.width = Cm(col_widths[j])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.paragraphs[0].paragraph_format.space_before = Pt(2)
            cell.paragraphs[0].paragraph_format.space_after = Pt(2)
    return table

def add_callout(doc, text, style='info'):
    # info / warning / success / danger
    colors = {
        'info':    (C_LIGHT,                RGBColor(0x1A, 0x52, 0x76)),
        'warning': (RGBColor(0xFF,0xF3,0xCD), C_ORANGE),
        'success': (RGBColor(0xD4,0xED,0xDA), C_GREEN),
        'danger':  (RGBColor(0xF8,0xD7,0xDA), C_RED),
    }
    bg, fg = colors.get(style, colors['info'])
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.cell(0, 0)
    set_cell_bg(cell, bg)
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.size = Pt(9.5)
    run.font.color.rgb = fg
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    doc.add_paragraph()  # spacing after callout

def add_spacer(doc, space=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space)
    p.paragraph_format.space_after = Pt(0)

# ════════════════════════════════════════════════════════════════════════════════
# 문서 생성 시작
# ════════════════════════════════════════════════════════════════════════════════
doc = Document()

# 기본 폰트 설정 (한글 호환)
style = doc.styles['Normal']
style.font.name = 'Malgun Gothic'
style.font.size = Pt(10)
style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')

# 여백 설정
section = doc.sections[0]
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(3.0)
section.right_margin  = Cm(2.5)

# ── 표지 타이틀 ─────────────────────────────────────────────────────────────
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_title.paragraph_format.space_before = Pt(20)
r = p_title.add_run("GHS 위험 라벨 수집 파이프라인")
r.font.size = Pt(22)
r.font.bold = True
r.font.color.rgb = C_NAVY
r.font.name = 'Malgun Gothic'
r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p_sub.add_run("현황 진단 및 개선 방향 보고서")
r2.font.size = Pt(15)
r2.font.color.rgb = C_BLUE
r2.font.name = 'Malgun Gothic'
r2._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')

p_date = doc.add_paragraph()
p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p_date.add_run("물성(Formulation) 팀 | 2026년 7월 14일")
r3.font.size = Pt(10)
r3.font.color.rgb = C_GRAY
p_date.paragraph_format.space_after = Pt(16)

doc.add_paragraph()

# ── 요약 (Executive Summary) ─────────────────────────────────────────────────
add_callout(doc,
    "【핵심 요약】  GHS 위험 라벨 예측 모델 개발을 위해 PDF 섹션 추출 → 데이터 풍부화 → "
    "라벨 계층화 → 품질 검증의 4단계 파이프라인을 구축·실행 완료하였다. "
    "섹션 매칭률은 39.6%(663/1,675행)로 데이터 상한선(40.0%) 직전까지 개선되었다. "
    "그러나 전체 데이터의 60%는 PDF 경로 자체가 없어 코드 개선만으로는 해결이 불가능하며, "
    "고신뢰 라벨(gold)은 291행에 그쳐 추가 데이터 수집이 모델 개발의 선결 과제이다.",
    style='info'
)

# ══════════════════════════════════════════════════════════════════════════════
# 1. 작업 개요 및 배경
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "1. 작업 개요 및 배경")

add_heading(doc, "1.1 팀 역할 분담", 2)
add_body(doc,
    "농약 제형(Formulation) 데이터에서 GHS 위험 분류 예측 모델을 개발하는 것이 물성팀의 목표이다. "
    "성분팀이 개별 성분의 SMILES·독성 데이터를 수집하는 반면, 물성팀은 제형 단위의 위험 라벨 "
    "(h_codes, ghs_hazard_classes, signal_word) 및 제형 특성(물리화학적 성질, 제형 코드 등)을 수집·정제한다."
)

add_heading(doc, "1.2 모델 개발 전략 (Two-Track)", 2)
add_section_table(doc,
    ["트랙", "방법론", "입력 데이터", "활용 조건"],
    [
        [("Track A", None), ("직접 예측 (Direct QSAR)", None),
         ("제형 특성 + 성분 SMILES", None), ("충분한 고신뢰 라벨 확보 시", None)],
        [("Track B", None), ("GHS CT Hybrid / 잔차 학습", None),
         ("성분별 GHS CT 임계값 + 제형 배합비", None), ("라벨 부족 시 선행 적용", None)],
    ],
    col_widths=[2.5, 4.5, 5.0, 4.0]
)
add_spacer(doc)
add_body(doc,
    "현 시점은 고신뢰 라벨이 291행으로 제한적이므로, "
    "Track B(GHS CT 가법성 가정)를 베이스라인으로 먼저 확립한 뒤 "
    "Track A와의 성능 비교를 통해 최종 전략을 결정하는 것이 권장된다."
)

# ══════════════════════════════════════════════════════════════════════════════
# 2. 파이프라인 구성 및 실행 결과
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "2. 파이프라인 구성 및 실행 결과")

add_heading(doc, "2.1 전체 파이프라인 흐름", 2)
add_section_table(doc,
    ["단계", "스크립트", "주요 작업", "출력"],
    [
        [("Phase 0.5", None), ("pdf_section_extractor.py", None),
         ("PDF/HTML → SDS 섹션별 JSON 추출 (LLM 없이 정규식)", None),
         ("section_extracts/*.json (490개)", None)],
        [("Phase 1", None), ("enrich_with_sections.py", None),
         ("master workbook + section JSON 조인 → 풍부화 CSV", None),
         ("artifacts/enriched_queue.csv (1,675행)", None)],
        [("Phase 2", None), ("extract_formulation_features.py", None),
         ("AWS Bedrock LLM 5개 에이전트 병렬 호출 → 제형 특성 추출", None),
         ("artifacts/formulation_characteristics_output.csv", None)],
        [("Step 3", None), ("(Phase 2 내)", None),
         ("label_tier 분류 (gold/silver/none)", None),
         ("label_tier 컬럼 추가", None)],
        [("Step 4", None), ("export_label_strata.py", None),
         ("gold/combined 서브셋 CSV 분리 출력", None),
         ("workflow/label_strata/*.csv", None)],
        [("Step 5", None), ("validate_labels.py", None),
         ("H-code 형식·signal_word 유효성 검증", None),
         ("workflow/diagnostics/label_validation.csv", None)],
    ],
    col_widths=[2.0, 5.5, 6.5, 5.0]
)
add_spacer(doc)

add_heading(doc, "2.2 이번 개선 사항 (금번 작업)", 2)
add_bullet(doc, "보조 디렉토리(ordered_option_runs/highlighted_sources/) 추가 처리: PDF 66개 + HTML 13개 = 79개 신규")
add_bullet(doc, "INPUT_COLS에 Supplemental_* 컬럼 2종 추가 → 이전 버전에서 조용히 누락되던 데이터 복구")
add_bullet(doc, "2-전략 source_id 매칭 구현: hex 부분 추출(전략 1) + 전체 stem 폴백(전략 2)")
add_bullet(doc, "파이프 구분 다중 경로 처리 (path1 | path2 형식 24행 대응)")
add_bullet(doc, "label_tier 시스템 도입: gold(pdf_regex 근거) / silver(값 있으나 근거 없음) / none")
add_bullet(doc, "라벨 검증 스크립트: H361d, H360Df 등 서브분류 변형 허용, CAUTION EPA 표준 추가")

# ══════════════════════════════════════════════════════════════════════════════
# 3. 핵심 지표 현황
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "3. 핵심 지표 현황")

add_heading(doc, "3.1 섹션 매칭률 (가장 중요한 지표)", 2)
add_kv_table(doc, [
    ("전체 제형 행",        "1,675행"),
    ("PDF 경로 있는 행",    "670행 (40.0%) ← 데이터 상한선",  "orange"),
    ("실제 매칭 달성",      "663행 (39.6%) ← 상한선까지 0.4%p",  "orange"),
    ("개선 전 매칭률",      "33.6% (562행)"),
    ("이번 개선폭",         "+6.0%p (+101행)",  "green"),
    ("PDF 경로 없는 행",    "1,005행 (60.0%) ← 코드로 해결 불가",  "red"),
])
add_spacer(doc)

add_callout(doc,
    "⚠ 병목 진단: 섹션 매칭률 33.6% → 39.6% 개선은 코드 수정(보조 디렉토리 추가, 컬럼 누락 수정)으로 "
    "달성 가능한 최대치에 근접한 것이다. 나머지 60%는 PDF 원본 파일이 없어 구조 추출이 근본적으로 불가능하다. "
    "이 1,005행에 대한 데이터 수집(PDF 확보 또는 외부 DB 매칭)이 다음 단계의 핵심 과제이다.",
    style='warning'
)

add_heading(doc, "3.2 GHS 라벨 품질 분포", 2)
add_section_table(doc,
    ["라벨 근거 유형", "ghs_hazard_classes", "h_codes", "의미"],
    [
        [("pdf_regex (Gold)", 'green'), ("~13.5% (약 226행)", 'green'),
         ("일부 포함", None), ("PDF 직접 추출 — 최고 신뢰도", None)],
        [("value_no_pdf (Silver)", 'orange'), ("~19.7% (약 330행)", 'orange'),
         ("일부 포함", None), ("LLM 추론값 — 노이즈 위험 있음", None)],
        [("none", 'red'), ("~65.8% (약 1,102행)", 'red'),
         ("대부분", None), ("값 자체 없음", None)],
    ],
    col_widths=[4.0, 4.0, 3.0, 7.0]
)
add_spacer(doc)

add_heading(doc, "3.3 label_tier 분포 (현재 기준)", 2)
add_kv_table(doc, [
    ("gold (최고 신뢰)",     "291행 — pdf_regex 근거 확인된 라벨",  "green"),
    ("combined (gold+silver)", "559행 — gold 포함 silver 확장 세트",  "orange"),
    ("none",                 "1,116행 — 라벨 없음 (모델 학습 제외 대상)",  "red"),
])
add_spacer(doc)

add_heading(doc, "3.4 라벨 검증 결과 (Step 5)", 2)
add_section_table(doc,
    ["검증 항목", "결과", "판정"],
    [
        [("H-code 형식 오류 (invalid_hcode)", None), ("0건", 'green'), ("✅ 합격", 'green')],
        [("signal_word 형식 오류 (invalid_signal_word)", None), ("0건", 'green'), ("✅ 합격", 'green')],
        [("gold 행 중 h_codes 비어있음 (gold_empty_hcodes)", None),
         ("49건 (gold 291행의 16.8%)", 'orange'),
         ("⚠ 확인 필요", 'orange')],
        [("gold 행 중 signal_word 무효값 (gold_null_signal_word)", None),
         ("11건 (NONE/UNKNOWN)", 'orange'),
         ("⚠ 확인 필요", 'orange')],
    ],
    col_widths=[7.5, 4.5, 3.0]
)
add_spacer(doc)
add_body(doc,
    "※ gold_empty_hcodes 49건은 ghs_hazard_classes에서 pdf_regex 근거가 확인된 행이므로 "
    "ghs_hazard_classes Y변수로는 유효하다. 단, h_codes를 별도 Y변수로 사용할 경우 이 49건은 제외해야 한다."
)

# ══════════════════════════════════════════════════════════════════════════════
# 4. 진단: 막힌 지점 상세 분석
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "4. 데이터 수집 병목 상세 진단")

add_heading(doc, "4.1 병목 1 — PDF 원본 부재 (최우선 해결 과제)", 2)
add_kv_table(doc, [
    ("영향 행 수",  "1,005행 (전체의 60.0%)",  "red"),
    ("원인",        "formulation_ingredients_master_audited.xlsx에 Audit_Source_File, "
                    "Supplemental_Source_File 모두 공란"),
    ("증상",        "section_extracts JSON 생성 자체 불가 — 어떤 코드 수정으로도 해결 불가"),
    ("필요 조치",   "해당 제형의 SDS/라벨 PDF 원본 확보 또는 외부 DB(EPA, ECHA, 국내 농약 DB) 매칭",  "orange"),
])
add_spacer(doc)

add_heading(doc, "4.2 병목 2 — 섹션 미탐지율 (섹션 추출 실패)", 2)
add_kv_table(doc, [
    ("영향 행 수",  "663행 중 약 ~80행 (추정) — section_2 미추출"),
    ("원인",
     "PDF 내 섹션 헤더 표기가 비표준 (예: 'HAZARD SUMMARY' 대신 '2. 위험성 개요' 한국어 표기, "
     "또는 연속 텍스트 형식 라벨)"),
    ("현재 대응",
     "label_header fallback 처리(섹션 없는 경우 전문 앞 2500자에서 직접 추출)로 부분 보완"),
    ("추가 가능 개선",
     "SECTION_STARTS 패턴에 한국어 패턴 추가 및 화학물질안전원 SDS 양식 대응",  "orange"),
])
add_spacer(doc)

add_heading(doc, "4.3 병목 3 — LLM 추론 라벨 신뢰도 (Silver 라벨 문제)", 2)
add_body(doc,
    "value_no_pdf 라벨(약 330행)은 PDF 근거 없이 LLM이 제형명·성분명만 보고 추론한 값이다. "
    "이 라벨을 학습에 포함하면 모델이 '진짜 위험 신호'가 아니라 'LLM의 추론 패턴'을 학습할 위험이 있다."
)
add_bullet(doc, "ghs_hazard_classes_evidence == value_no_pdf: ~19.7% (약 330행)")
add_bullet(doc, "이 중 실제 정답과 일치하는 비율 미검증 → 무작위 샘플 50건 수동 검증 필요")
add_bullet(doc, "권장: 초기 모델은 gold(291행)만으로 baseline 수립 후 silver 추가 효과 비교")
add_spacer(doc)

add_heading(doc, "4.4 병목 4 — 타겟 라벨 불균형", 2)
add_section_table(doc,
    ["Y 변수", "유효값 행 수 (추정)", "gold만", "문제점"],
    [
        [("ghs_hazard_classes", None), ("~556행 (33.2%)", 'orange'), ("~226행", None),
         ("클래스 수 다수, 멀티라벨 학습 필요", None)],
        [("h_codes", None), ("~405행 (24.1%)", 'red'), ("~242행", None),
         ("gold 중 49행이 h_codes 없음", None)],
        [("signal_word", None), ("~595행 (35.5%)", 'orange'), ("일부", None),
         ("3-class 분류(DANGER/WARNING/CAUTION)", None)],
    ],
    col_widths=[4.0, 4.5, 3.0, 6.5]
)
add_spacer(doc)
add_callout(doc,
    "signal_word (3-class)가 가장 높은 유효값 비율(35.5%)을 보이므로 "
    "초기 모델의 Y변수로 가장 현실적이다. ghs_hazard_classes는 멀티라벨로 구성이 복잡하지만 "
    "최종 목표 변수이므로 병행 수집해야 한다.",
    style='info'
)

# ══════════════════════════════════════════════════════════════════════════════
# 5. 보완 및 개선 방향
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "5. 보완 및 개선 방향")

add_heading(doc, "5.1 즉시 실행 가능 (코드/파이프라인 개선)", 2)
add_section_table(doc,
    ["우선순위", "작업", "예상 효과", "담당"],
    [
        [("★★★", 'red'), ("Phase 2 재실행 (AWS Bedrock)", None),
         ("새 enriched_queue.csv(663행) 반영 → gold 291 → 350+ 예상", None),
         ("물성팀", None)],
        [("★★☆", 'orange'), ("한국어 섹션 헤더 패턴 추가 (pdf_section_extractor.py)", None),
         ("국내 SDS 섹션 미탐지 보완, section2 추출률 향상", None),
         ("물성팀", None)],
        [("★★☆", 'orange'), ("Silver 라벨 50건 무작위 수동 검증", None),
         ("value_no_pdf 라벨의 실제 정확도 파악 → 학습 포함 여부 결정", None),
         ("물성팀", None)],
        [("★☆☆", None), ("gold_empty_hcodes 49건 수동 보완", None),
         ("h_codes Y변수 gold 세트 완성도 향상", None),
         ("물성팀", None)],
    ],
    col_widths=[2.2, 6.5, 5.0, 2.3]
)
add_spacer(doc)

add_heading(doc, "5.2 데이터 수집 우선 과제 (중기)", 2)
add_bullet(doc, "【최우선】 PDF 미확보 1,005행 대상 SDS/라벨 원본 수집", bold=True, color=C_RED)
add_bullet(doc,
    "우선 순위: EPA 등록 번호 있는 행 → EPA PPIS/CDX 공개 데이터베이스에서 라벨 PDF 직접 다운로드",
    level=1
)
add_bullet(doc,
    "ECHA C&L Inventory: EU 농약 제형의 GHS 분류 공개 데이터 활용 가능",
    level=1
)
add_bullet(doc,
    "국내: 농약 허가 정보시스템(농촌진흥청), 화학물질안전원 MSDS DB",
    level=1
)
add_spacer(doc, 2)
add_bullet(doc, "pH, 계면활성제 정보 수집 (제형 특성 핵심 변수)", bold=True, color=C_ORANGE)
add_bullet(doc, "현재 파이프라인에서 pH가 누락됨 — 안구/피부 자극성 예측에 중요 피처", level=1)
add_bullet(doc, "SDS Section 9 (물리화학적 성질)에서 추출 필요 — 현재 Section 1, 2, 14만 처리 중", level=1)
add_spacer(doc, 2)
add_bullet(doc, "성분 SMILES 완성도 향상 (성분팀 협력)", bold=True)
add_bullet(doc, "Track A 모델의 핵심 입력값 — 성분팀 데이터와 JOIN 기준점 확립 필요", level=1)

add_heading(doc, "5.3 섹션 추출 고도화 (pdf_section_extractor.py)", 2)
add_section_table(doc,
    ["개선 항목", "현재 상태", "개선 방안"],
    [
        [("Section 9 추출", None), ("미구현", 'red'),
         ("pH, 물리적 상태, 밀도 등 제형 특성 변수 추가 추출", None)],
        [("한국어 SDS 패턴", None), ("영문 패턴만 존재", 'orange'),
         ("'제2절', '유해·위험성' 등 화학물질안전원 표준 양식 대응", None)],
        [("GHS 심볼 파싱", None), ("텍스트만 파싱", None),
         ("HTML SDS에서 GHS 픽토그램 alt text 추출 추가", None)],
        [("비구조화 라벨 처리", None), ("label_header fallback만", None),
         ("미국 EPA RTF/라벨 형식의 ACTIVE INGREDIENT 섹션 특화 파싱", None)],
    ],
    col_widths=[3.5, 4.5, 10.0]
)
add_spacer(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 6. 모델 개발 로드맵
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "6. 모델 개발 로드맵 (데이터 우선)")

add_section_table(doc,
    ["단계", "기간 (예상)", "목표", "선결 조건"],
    [
        [("Phase M0\n데이터 기반 확립", None),
         ("현재 ~ 2026.08", 'orange'),
         ("Gold 라벨 500행 이상 확보\nPDF 원본 300건 추가 수집\nPhase 2 재실행 → label_tier 갱신", None),
         ("PDF 추가 수집\nAWS Bedrock 실행", None)],
        [("Phase M1\nBaseline 모델", None),
         ("2026.08 ~ 09", None),
         ("Track B: GHS CT 가법성 베이스라인\nsignal_word 3-class 분류기 (gold 291행)\n성능 지표 확립 (AUC, F1 macro)", None),
         ("Gold ≥291행 확정\n성분 SMILES 연결", None)],
        [("Phase M2\nTrack A 실험", None),
         ("2026.09 ~ 10", None),
         ("SMILES 기반 직접 예측 모델 구축\nghs_hazard_classes 멀티라벨 분류\nTrack A vs B 성능 비교", None),
         ("Gold ≥500행\nSMILES 커버리지 ≥60%", None)],
        [("Phase M3\n모델 고도화", None),
         ("2026.10 ~", None),
         ("Silver 라벨 포함 효과 검증\n앙상블 / Transfer Learning 실험\n불확실성 추정 추가", None),
         ("Silver 검증 완료\n모델 성능 기준선 확립", None)],
    ],
    col_widths=[3.0, 3.0, 8.0, 4.0]
)
add_spacer(doc)

add_callout(doc,
    "【현 시점 판단】 모델 개발보다 데이터 수집이 명확한 선결 과제이다. "
    "Gold 291행은 5-fold CV에서 각 fold 당 ~58행으로, 고차원 피처를 가진 멀티라벨 분류에는 "
    "과소적합 위험이 높다. PDF 원본 300건 추가 확보 및 Phase 2 재실행으로 gold를 "
    "500행 수준으로 올린 뒤 Track B 베이스라인을 우선 수립하는 것을 권장한다.",
    style='warning'
)

# ══════════════════════════════════════════════════════════════════════════════
# 7. 검증 기준 및 수락 조건
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "7. 검증 기준 및 데이터 품질 수락 조건")

add_section_table(doc,
    ["항목", "현재 값", "목표 기준", "판정"],
    [
        [("섹션 매칭률", None), ("39.6%", None), ("≥38% (달성)", None), ("✅ 달성", 'green')],
        [("H-code 형식 오류", None), ("0건", None), ("0건", None), ("✅ 합격", 'green')],
        [("signal_word 형식 오류", None), ("0건", None), ("0건", None), ("✅ 합격", 'green')],
        [("label_tier null", None), ("0건", None), ("0건", None), ("✅ 합격", 'green')],
        [("gold 행 수", None), ("291행", None), ("≥500행 (모델 개발용)", None), ("⚠ 미달", 'red')],
        [("gold ghs_null 수", None), ("확인 필요", None), ("0 (목표)", None), ("⏳ 확인 중", 'orange')],
        [("Silver 검증 정확도", None), ("미검증", None), ("≥80% 일치율", None), ("⏳ 미실행", 'orange')],
        [("Section 9 추출률", None), ("0% (미구현)", None), ("≥30%", None), ("🔴 미구현", 'red')],
    ],
    col_widths=[4.5, 3.0, 4.5, 3.0]
)
add_spacer(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 8. 다음 단계 액션 아이템
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "8. 다음 단계 액션 아이템")

add_section_table(doc,
    ["#", "액션", "우선순위", "담당", "비고"],
    [
        [("#1", 'red'), ("Phase 2 재실행 (extract_formulation_features.py)", None),
         ("즉시", 'red'), ("물성팀", None), ("AWS Bedrock 인증 필요", None)],
        [("#2", 'red'), ("Step 4, 5 재실행 → 갱신된 gold 수 확인", None),
         ("즉시", 'red'), ("물성팀", None), ("#1 완료 후", None)],
        [("#3", 'orange'), ("Silver 라벨 50건 무작위 수동 검증", None),
         ("1주 내", 'orange'), ("물성팀", None), ("샘플링 스크립트 작성 필요", None)],
        [("#4", 'orange'), ("PDF 미확보 1,005행 EPA Reg No. 추출 → EPA DB 조회", None),
         ("1주 내", 'orange'), ("물성팀", None), ("EPA PPIS 공개 API 활용", None)],
        [("#5", None), ("pdf_section_extractor.py에 Section 9 추출 로직 추가", None),
         ("2주 내", None), ("물성팀", None), ("pH, 물리적 상태 등", None)],
        [("#6", None), ("한국어 SDS 섹션 패턴 추가", None),
         ("2주 내", None), ("물성팀", None), ("농촌진흥청 MSDS 양식 기준", None)],
        [("#7", None), ("성분팀과 SMILES 커버리지 현황 공유 및 JOIN 기준 협의", None),
         ("이번 주", None), ("물성+성분팀", None), ("Track A 선결 조건", None)],
    ],
    col_widths=[1.0, 6.5, 2.0, 2.0, 4.5]
)
add_spacer(doc)

# 푸터
doc.add_paragraph()
p_foot = doc.add_paragraph()
p_foot.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r_foot = p_foot.add_run("물성(Formulation) 팀 | 작성일: 2026-07-14 | 내부 배포용")
r_foot.font.size = Pt(8)
r_foot.font.color.rgb = C_GRAY
r_foot.font.italic = True

doc.save(OUT_PATH)
print(f"✅ 보고서 저장 완료: {OUT_PATH}")
